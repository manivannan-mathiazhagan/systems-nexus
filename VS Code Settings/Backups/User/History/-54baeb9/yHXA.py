# ****************************************************************************************************
# Script Name    : Veristat_TLF_Packager.py
#
# Purpose        : GUI Tool to scan RTF, DOCX, and PDF files in a folder,
#                  extract TLF titles (from RTF: header/first lines, DOCX: header/body, PDF: page 1),
#                  display for user review, filter, reorder (move up/down), and export details to Excel.
#                  Users can convert RTF/DOCX files to PDF (via Microsoft Word),
#                  and generate a single bookmarked PDF with advanced, clickable Table of Contents.
#
# Author         : Manivannan Mathialagan
#
# Created On     : 22-May-2025
#
# Key Parameters (GUI-based):
#   - Folder Selection         : User selects the folder containing RTF, DOCX, and PDF files.
#   - Output PDF Name          : User specifies the name for the merged PDF.
#
# Key Features:
#   - Extracts and displays TLF titles from headers/bodies of RTF, DOCX, and first page of PDF.
#   - Allows user to filter, select, and reorder outputs for packaging.
#   - Bookmark/Include columns are interactive; Bookmark is editable.
#   - Exports details and custom order to Excel for documentation or further review.
#   - Converts RTF/DOCX to PDF using Microsoft Word (requires pywin32).
#   - Generates a merged PDF with bookmarks and advanced, wrapped, clickable Table of Contents (TOC).
#
# Example Usage:
#   - Run the script: python TLF_Packager.py
#   - Use the GUI to select folder, review/reorder files, and generate outputs.
#
# Notes:
#   - Requires Python 3, openpyxl, PyMuPDF (fitz), PyQt5, pywin32, python-docx
#   - Requires Microsoft Word installed for RTF/DOCX to PDF conversion (via pywin32)
#   - For any issues or suggestions, contact manivannan.mathialagan@veristat.com
# ****************************************************************************************************

import sys
import subprocess
import site

def ensure_user_site():
    user_site = site.getusersitepackages()
    if user_site not in sys.path:
        sys.path.insert(0, user_site)

def install_if_missing(package_name, import_name=None):
    try:
        __import__(import_name or package_name)
    except ImportError:
        print(f"[Installing] {package_name} (user mode)...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", package_name])
        print(f"[Done] {package_name} installed.")

ensure_user_site()
install_if_missing("openpyxl")
install_if_missing("PyMuPDF", "fitz")
install_if_missing("PyQt5")
install_if_missing("pywin32")
install_if_missing("python-docx", "docx")

import os
import re
import fitz
import openpyxl
import docx
import datetime
import time
from openpyxl import Workbook
from PyQt5 import QtCore, QtGui, QtWidgets

CHECKED = '\u2611'   # ☑
UNCHECKED = '\u2610' # ☐

# --- Extraction and helper functions ---

def extract_groups_from_rtf_header(rtf_text):
    header_match = re.search(r'(\\header[lr]?)(.+?)(\\sectd|$)', rtf_text, re.DOTALL)
    if not header_match: return []
    header_text = header_match.group(2)
    groups, depth, buf = [], 0, ""
    for c in header_text:
        if c == '{':
            if depth == 0: buf = ""
            depth += 1
            buf += c
        elif c == '}':
            buf += c; depth -= 1
            if depth == 0 and buf: groups.append(buf); buf = ""
        elif depth > 0:
            buf += c
    return groups

def rtf_group_to_text(group):
    group = group.replace('\\cell', '\n').replace('\\row', '\n')
    text = re.sub(r'\\[a-z]+\d*[\s]?|{|}', '', group)
    text = re.sub(r'(cellx\d+|x\d+)', '', text)
    text = re.sub(r'\n+', '\n', text)
    lines = [re.sub(r'^[-: ]+|[-: ]+$', '', l.strip()) for l in text.splitlines() if l.strip()]
    return lines

def extract_header_lines_with_tablepattern(rtf_path):
    try:
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_text = f.read()
        groups = extract_groups_from_rtf_header(rtf_text)
        lines = []
        for group in groups:
            lines += rtf_group_to_text(group)
        return lines
    except Exception:
        return []

def valid_title_candidate(candidate):
    c = candidate.strip()
    if not c: return False
    if c.lower().startswith("note:"): return False
    if c.count('.') > 1: return False
    if re.search(r'\[\d+\]', c): return False
    if re.match(r'^\\?\*.*$', c): return False
    if re.match(r'^[A-Z0-9\\\*]+$', c): return False
    return True

def filter_note_from_line(text):
    return re.split(r'\bNote\s*:', text, flags=re.IGNORECASE)[0].strip()

def extract_titles_from_rtf(rtf_path):
    header_lines = extract_header_lines_with_tablepattern(rtf_path)
    title1, title2, title3 = "", "", ""
    bookmark = ""
    title_pattern = r'(?:(Table|Listing|Figure)\s+)?(Appendix)\s+\d+(?:\.\d+)*[A-Za-z0-9]*|(Table|Listing|Figure)\s+\d+(?:\.\d+)*[A-Za-z0-9]*'
    for i, line in enumerate(header_lines):
        match = re.search(title_pattern, line, re.IGNORECASE)
        if match:
            title1 = match.group(0).strip()
            remainder = line[len(match.group(0)):].strip(" :–-")
            remainder = filter_note_from_line(remainder)
            if remainder and valid_title_candidate(remainder): title2 = remainder
            if not title2 and i+1 < len(header_lines):
                t2 = header_lines[i+1]; t2_clean = filter_note_from_line(t2)
                if valid_title_candidate(t2_clean): title2 = t2_clean.strip()
            if i+2 < len(header_lines):
                t3 = header_lines[i+2]; t3_clean = filter_note_from_line(t3)
                if valid_title_candidate(t3_clean): title3 = t3_clean.strip()
            break
    if title1:
        bookmark = title1
        if title2: bookmark += ": " + title2
        if title3: bookmark += " - " + title3
        return title1, title2, title3, bookmark
    return "", "", "", os.path.basename(rtf_path)

def extract_title_from_pdf(pdf_path):
    title_pattern = r'(?:(Table|Listing|Figure)\s+)?(Appendix)\s+\d+(?:\.\d+)*[A-Za-z0-9]*|(Table|Listing|Figure)\s+\d+(?:\.\d+)*[A-Za-z0-9]*'
    try:
        doc = fitz.open(pdf_path)
        text = doc[0].get_text("text")
        doc.close()
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        title1, title2, title3 = "", "", ""
        for i, line in enumerate(lines[:20]):
            match = re.search(title_pattern, line, re.IGNORECASE)
            if match:
                title1 = match.group(0).strip()
                remainder = line[len(match.group(0)):].strip(" :–-")
                remainder = filter_note_from_line(remainder)
                if remainder and valid_title_candidate(remainder): title2 = remainder
                if not title2 and i+1 < len(lines):
                    cand2 = lines[i+1]; cand2_clean = filter_note_from_line(cand2)
                    if valid_title_candidate(cand2_clean): title2 = cand2_clean.strip()
                if i+2 < len(lines):
                    cand3 = lines[i+2]; cand3_clean = filter_note_from_line(cand3)
                    if valid_title_candidate(cand3_clean): title3 = cand3_clean.strip()
                break
        if not title1:
            for idx, line in enumerate(lines[:3]):
                if valid_title_candidate(line):
                    if not title1: title1 = filter_note_from_line(line)
                    elif not title2: title2 = filter_note_from_line(line)
                    elif not title3: title3 = filter_note_from_line(line)
        bookmark = title1
        if title2: bookmark += ": " + title2
        if title3: bookmark += " - " + title3
        if not bookmark.strip():
            bookmark = os.path.basename(pdf_path)
        return title1, title2, title3, bookmark
    except Exception:
        return "", "", "", os.path.basename(pdf_path)

def extract_title_from_docx(docx_path):
    title_pattern = r'(?:(Table|Listing|Figure)\s+)?(Appendix)\s+\d+(?:\.\d+)*[A-Za-z0-9]*|(Table|Listing|Figure)\s+\d+(?:\.\d+)*[A-Za-z0-9]*'
    try:
        doc = docx.Document(docx_path)
        paras = []
        for section in doc.sections:
            for h_para in section.header.paragraphs:
                line = h_para.text.strip()
                if line:
                    paras.append(line)
        title1 = title2 = title3 = ""
        for idx, line in enumerate(paras[:10]):
            match = re.match(title_pattern, line, re.IGNORECASE)
            if match:
                title1 = match.group(0).strip()
                remainder = line[len(match.group(0)):].strip(" :–-")
                remainder = filter_note_from_line(remainder)
                if remainder and valid_title_candidate(remainder): title2 = remainder
                if not title2 and idx+1 < len(paras):
                    cand2 = paras[idx+1]
                    if valid_title_candidate(cand2): title2 = cand2
                if idx+2 < len(paras):
                    cand3 = paras[idx+2]
                    if valid_title_candidate(cand3): title3 = cand3
                break
        if not title1:
            body_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text: body_paras.append(cell_text)
            for idx, line in enumerate(body_paras[:15]):
                match = re.match(title_pattern, line, re.IGNORECASE)
                if match:
                    title1 = match.group().strip()
                    remainder = line[len(match.group()):].strip(" :–-")
                    remainder = filter_note_from_line(remainder)
                    if remainder and valid_title_candidate(remainder): title2 = remainder
                    if not title2 and idx+1 < len(body_paras):
                        cand2 = body_paras[idx+1]
                        if valid_title_candidate(cand2): title2 = cand2
                    if idx+2 < len(body_paras):
                        cand3 = body_paras[idx+2]
                        if valid_title_candidate(cand3): title3 = cand3
                    break
            if not title1:
                all_lines = paras + body_paras
                found_lines = 0
                for line in all_lines[:15]:
                    if valid_title_candidate(line):
                        found_lines += 1
                        if not title1: title1 = filter_note_from_line(line)
                        elif not title2: title2 = filter_note_from_line(line)
                        elif not title3: title3 = filter_note_from_line(line)
                        if found_lines == 3: break
        bookmark = title1
        if title2: bookmark += ": " + title2
        if title3: bookmark += " - " + title3
        if not bookmark.strip():
            bookmark = os.path.basename(docx_path)
        return title1, title2, title3, bookmark
    except Exception:
        return "", "", "", os.path.basename(docx_path)

def parse_sort_key(bookmark):
    order = {"table": 0, "listing": 1, "figure": 2}
    text = (bookmark or "").lower()
    text = re.sub(r'^[^a-zA-Z]*', '', text)
    if 'appendix' in text:
        match = re.search(r'appendix\s*(\d+)', text)
        num = [int(match.group(1))] if match else [9999]
        return (3, num)
    for key in order:
        if key in text:
            match = re.search(r'(\d+(\.\d+)*)', text)
            num = [int(x) for x in match.group(1).split('.')] if match else [0]
            return (order[key], num)
    return (99, [0])

def wrap_toc_title(title, font_size, max_width):
    words = (title or "").split()
    lines = []
    current_line = ""
    for word in words:
        test_line = current_line + (" " if current_line else "") + word
        if fitz.get_text_length(test_line, fontsize=font_size) <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return [line for line in lines if line.strip()]
def add_toc_and_links(bookmarked_pdf, final_pdf, font_size=8):
    def extract_toc_entries(doc):
        return [(lvl, title.strip(), page_num - 1) for lvl, title, page_num in doc.get_toc(simple=True)]

    def paginate_wrapped_entries(toc_entries, font_size, toc_text_max_width, page_height):
        y_spacing = font_size * 1.5
        lines_per_page = int((page_height - 100) // y_spacing) - 2
        paginated_entries, current_page_entries, current_line_count = [], [], 0
        for entry in toc_entries:
            level, title, target_page = entry
            indent = 20 * (level - 1)
            entry_max_width = toc_text_max_width - indent
            wrapped_lines = wrap_toc_title(title, font_size, entry_max_width)
            line_count = len(wrapped_lines)
            if current_line_count + line_count > lines_per_page:
                paginated_entries.append(current_page_entries)
                current_page_entries, current_line_count = [], 0
            current_page_entries.append((entry, wrapped_lines))
            current_line_count += line_count
        if current_page_entries:
            paginated_entries.append(current_page_entries)
        return paginated_entries

    def generate_toc_pages(paginated_entries, font_size, page_width, page_height, toc_page_count,
                           toc_text_max_width, page_number_x, page_number_width, gap, right_margin):
        toc_doc = fitz.open()
        link_targets = []
        left_margin, top_margin = 50, 50
        y_spacing = font_size * 1.5
        for page_index, entries in enumerate(paginated_entries):
            page = toc_doc.new_page(width=page_width, height=page_height)
            y = top_margin
            if page_index == 0:
                toc_title = "Table of Contents"
                title_width = fitz.get_text_length(toc_title, fontsize=font_size + 2)
                page.insert_text(
                    (page_width / 2 - title_width / 2, y),
                    toc_title,
                    fontsize=font_size + 2,
                    fontname="Helvetica",
                    color=(0, 0, 0)
                )
                y += font_size * 2
            for (level, title, target_page), wrapped_lines in entries:
                indent = 20 * (level - 1)
                x = left_margin + indent
                page_number_str = str(target_page + toc_page_count + 1)
                page_number_width_local = fitz.get_text_length(page_number_str, fontsize=font_size)
                first_line_y = y
                for i, line in enumerate(wrapped_lines):
                    line_width = fitz.get_text_length(line, fontsize=font_size)
                    if i == len(wrapped_lines) - 1:
                        page.insert_text((x, y), line, fontsize=font_size)
                        dots_start_x = x + line_width + 2
                        if dots_start_x < page_number_x - page_number_width_local - gap:
                            dots_width = page_number_x - page_number_width_local - gap - dots_start_x
                            dot_char_width = fitz.get_text_length('.', fontsize=font_size)
                            dot_count = int(dots_width // dot_char_width)
                            page.insert_text((dots_start_x, y), '.' * dot_count, fontsize=font_size)
                        page.insert_text((page_number_x - page_number_width_local, y),
                                         page_number_str, fontsize=font_size)
                    else:
                        page.insert_text((x, y), line, fontsize=font_size)
                    y += y_spacing
                rect = fitz.Rect(x, first_line_y - font_size, page_number_x, y)
                link_targets.append((page_index, rect, target_page))
        return toc_doc, link_targets

    base = fitz.open(bookmarked_pdf)
    original_toc = extract_toc_entries(base)
    original_bookmarks = base.get_toc()
    width, height = fitz.paper_size("a4")
    left_margin, right_margin = 50, 60
    max_page_number = 99999
    page_number_width = fitz.get_text_length(str(max_page_number), fontsize=font_size)
    page_number_x = width - right_margin
    gap = 8
    toc_text_max_width = page_number_x - left_margin - page_number_width - gap

    paginated = paginate_wrapped_entries(original_toc, font_size, toc_text_max_width, height)
    toc_doc, link_targets = generate_toc_pages(
        paginated, font_size, width, height, len(paginated),
        toc_text_max_width, page_number_x, page_number_width, gap, right_margin
    )
    toc_page_count = len(paginated)

    final = fitz.open()
    final.insert_pdf(toc_doc)
    final.insert_pdf(base)

    for toc_page_index, rect, target_page in link_targets:
        final[toc_page_index].insert_link({
            "kind": fitz.LINK_GOTO,
            "from": rect,
            "page": target_page + toc_page_count
        })

    def shift_bookmark_pages(bookmarks, offset):
        return [[lvl, title, page + offset] for lvl, title, page in bookmarks if len([lvl, title, page]) >= 3]

    final.set_toc(shift_bookmark_pages(original_bookmarks, toc_page_count))
    final.save(final_pdf)
    final.close()
    base.close()

def export_bookmarks_to_excel(table_data, out_path):
    wb = Workbook()
    ws = wb.active
    headers = ["Include", "File Type", "Filename", "Title 1", "Title 2", "Title 3", "Bookmark", "Order"]
    ws.append(headers)
    for i, row in enumerate(table_data):
        ws.append([(1 if row[0] else 0)] + list(row[1:]) + [i+1])
    for idx, width in enumerate([8, 12, 30, 30, 30, 30, 50, 8], start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = width
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical='top')
    wb.save(out_path)

def convert_to_pdf_with_word(input_path, output_folder, logger=None):
    import win32com.client
    pdf_path = os.path.splitext(input_path)[0] + ".pdf"
    pdf_path = os.path.join(output_folder, os.path.basename(pdf_path))
    word = win32com.client.Dispatch('Word.Application')
    try:
        doc = word.Documents.Open(os.path.abspath(input_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)
        time.sleep(0.5)
    except Exception as e:
        pass
    finally:
        word.Quit()
    return pdf_path
class TLFApp(QtWidgets.QWidget):
    log_signal = QtCore.pyqtSignal(str)
    message_signal = QtCore.pyqtSignal(str, str)
    AUTO_REFRESH_INTERVAL_MS = 3000

    def __init__(self):
        super().__init__()
        icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "app.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QtGui.QIcon(icon_path))

        self.setWindowTitle("Veristat TLF Packager")
        self.setStyleSheet("background-color: #f4f4f4; border-radius: 14px;")
        self.resize(1280, 720)
        font_main = QtGui.QFont("Times New Roman", 12)
        font_header = QtGui.QFont("Times New Roman", 21, QtGui.QFont.Bold)
        font_table = QtGui.QFont("Times New Roman", 10)

        self.table_data = []
        self.out_folder = ""
        self.filtered_indices = []
        self.select_all_state = True  # Toggle tracker

        self.sorted_once = False
        self.current_sort_column = None
        self.sort_order = QtCore.Qt.AscendingOrder

        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)

        header = QtWidgets.QLabel("Veristat TLF Packager")
        header.setFont(font_header)
        header.setAlignment(QtCore.Qt.AlignCenter)
        header.setStyleSheet("color: black; background: #b3e3e7; padding: 14px 0 10px 0; border-radius: 18px;")
        layout.addWidget(header)

        query = QtWidgets.QLabel(
            "• For any queries, issues faced, or if you have ideas for improvement,\n  please contact Manivannan (manivannan.mathialagan@veristat.com)"
        )
        query.setFont(font_main)
        query.setStyleSheet("background: #f4f4f4; color: #222;")
        query.setAlignment(QtCore.Qt.AlignRight)
        layout.addWidget(query)

        # === Filter Row with TOC checkbox ===
        filter_layout = QtWidgets.QHBoxLayout()
        filter_label = QtWidgets.QLabel("Search:")
        filter_label.setFont(font_main)
        filter_layout.addWidget(filter_label)

        self.search_box = QtWidgets.QLineEdit()
        self.search_box.setFont(font_table)
        self.search_box.setPlaceholderText("Type to filter rows in any column...")
        self.search_box.setToolTip("Type to filter rows in any column (e.g. Table, RTF, T14.1.1, etc.)")
        self.search_box.setMaximumWidth(320)
        self.search_box.setStyleSheet("""
            QLineEdit { border-radius: 8px; padding: 4px 6px; border: 1px solid #b8c2d0; background: #fff; }
        """)
        filter_layout.addWidget(self.search_box)

        # === TOC Checkbox next to Search ===
        self.toc_checkbox = QtWidgets.QCheckBox("Include TOC page")
        self.toc_checkbox.setChecked(True)
        self.toc_checkbox.setFont(QtGui.QFont("Times New Roman", 11))
        filter_layout.addWidget(self.toc_checkbox)

        tip_label = QtWidgets.QLabel("Bookmark text is editable (double-click), and Include column can be checked/unchecked.")
        tip_label.setFont(QtGui.QFont("Times New Roman", 11))
        tip_label.setStyleSheet("background: #e6f2fb; color: #12608d; border-radius: 5px; padding: 3px 12px; margin-left: 14px;")
        filter_layout.addWidget(tip_label)
        filter_layout.addStretch()
        layout.addLayout(filter_layout)
        # === Table Headers ===
        self.labels = ["Include", "Type", "Filename", "Title 1", "Title 2", "Title 3", "Bookmark"]
        self.table = QtWidgets.QTableWidget(0, len(self.labels))
        self.table.setHorizontalHeaderLabels(self.labels)
        self.table.setFont(font_table)
        self.table.setEditTriggers(QtWidgets.QTableWidget.NoEditTriggers)
        self.table.setSelectionBehavior(QtWidgets.QTableWidget.SelectRows)
        self.table.setSelectionMode(QtWidgets.QTableWidget.SingleSelection)
        self.table.setShowGrid(True)
        self.table.setAlternatingRowColors(True)
        self.table.setStyleSheet("""
            QHeaderView::section { background-color: #b8c2d0; color: #222; font-weight: bold; border-radius: 10px; padding: 3px;}
            QTableWidget { border-radius: 14px; border: 1px solid #b8c2d0; background: #f8fbff; }
            QTableWidget::item:selected { background: #356AC3; color: #fff; }
        """)
        self.table.verticalHeader().setDefaultSectionSize(24)
        header_view = self.table.horizontalHeader()
        header_view.setSectionResizeMode(QtWidgets.QHeaderView.Stretch)
        self.table.setColumnWidth(0, 60)
        self.table.setColumnWidth(1, 75)
        layout.addWidget(self.table)

        # === Button Row ===
        btn_layout = QtWidgets.QHBoxLayout()
        btn_base_style = """
            QPushButton {
                font-family: 'Times New Roman';
                font-weight: bold;
                font-size: 11pt;
                border-radius: 10px;
                padding: 6px 16px;
            }
            QPushButton:disabled {
                background-color: #e1e1e1; color: #aaa;
            }
        """

        self.browse_btn = QtWidgets.QPushButton("Browse Folder")
        self.browse_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #b3c6e7; color: #222; }
            QPushButton:hover { background-color: #d5deef; }
        """)
        btn_layout.addWidget(self.browse_btn)

        # ✅ Combined Toggle Button
        self.select_toggle_btn = QtWidgets.QPushButton("Select All")
        self.select_toggle_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #33b9b2; color: #fff; }
            QPushButton:hover { background-color: #4eddd7; }
        """)
        btn_layout.addWidget(self.select_toggle_btn)

        self.export_btn = QtWidgets.QPushButton("Export to Excel")
        self.export_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #71c95e; color: #fff; }
            QPushButton:hover { background-color: #98eb84; }
        """)
        btn_layout.addWidget(self.export_btn)

        self.up_btn = QtWidgets.QPushButton("Move Up")
        self.up_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #ad7dfc; color: #fff; }
            QPushButton:hover { background-color: #c7a2fd; }
        """)
        btn_layout.addWidget(self.up_btn)

        self.down_btn = QtWidgets.QPushButton("Move Down")
        self.down_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #ad7dfc; color: #fff; }
            QPushButton:hover { background-color: #c7a2fd; }
        """)
        btn_layout.addWidget(self.down_btn)

        self.convert_btn = QtWidgets.QPushButton("Convert to PDF")
        self.convert_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #ff9933; color: #222; }
            QPushButton:hover { background-color: #ffbc80; }
        """)
        btn_layout.addWidget(self.convert_btn)

        self.pack_btn = QtWidgets.QPushButton("Pack && TOC")
        self.pack_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #89c990; color: #222; font-size: 13pt;}
            QPushButton:hover { background-color: #b2ebbb; }
        """)
        btn_layout.addWidget(self.pack_btn)

        # 🟥 Terminate button at far right
        self.terminate_btn = QtWidgets.QPushButton("Terminate")
        self.terminate_btn.setStyleSheet(btn_base_style + """
            QPushButton { background-color: #f14b4b; color: white; }
            QPushButton:hover { background-color: #f98a8a; }
        """)
        btn_layout.addWidget(self.terminate_btn)

        layout.addLayout(btn_layout)
        # === Button Tooltips ===
        self.browse_btn.setToolTip("Browse and select the folder containing your TLF outputs.")
        self.select_toggle_btn.setToolTip("Click to select or deselect all visible rows.")
        self.export_btn.setToolTip("Export the current table (with bookmarks) to an Excel file.")
        self.up_btn.setToolTip("Move selected row up in the order.")
        self.down_btn.setToolTip("Move selected row down in the order.")
        self.convert_btn.setToolTip("Convert selected RTF/DOCX files to PDF using MS Word.")
        self.pack_btn.setToolTip("Merge all selected files into a final bookmarked PDF.\nIncludes a Table of Contents if enabled.")
        self.terminate_btn.setToolTip("Force quit the application immediately.")

        # === Output File Field ===
        self.outpdf_var = QtWidgets.QLineEdit(f"TLFs_Merged_{datetime.datetime.now().strftime('%Y%m%d_T%H%M')}.pdf")
        self.outpdf_var.setStyleSheet("""
            QLineEdit { border-radius: 8px; padding: 4px 6px; border: 1px solid #b8c2d0; background: #fff; }
        """)
        outpdf_lbl = QtWidgets.QLabel("Output PDF Name:")
        outpdf_lbl.setFont(font_main)
        output_pdf_row = QtWidgets.QHBoxLayout()
        output_pdf_row.addWidget(outpdf_lbl)
        output_pdf_row.addWidget(self.outpdf_var)
        layout.addLayout(output_pdf_row)

        # === Log Box ===
        log_lbl = QtWidgets.QLabel("Log:")
        log_lbl.setFont(font_main)
        layout.addWidget(log_lbl)
        self.log_box = QtWidgets.QPlainTextEdit()
        self.log_box.setReadOnly(True)
        self.log_box.setMaximumHeight(130)
        self.log_box.setFont(QtGui.QFont("Times New Roman", 10))
        self.log_box.setStyleSheet("background: #f7f7f7; color: #1a2a4f; border-radius: 8px;")
        layout.addWidget(self.log_box)

        self.log_signal.connect(self.log)
        self.message_signal.connect(self.show_messagebox)

        self.set_button_states(browse=True, others=False)
        self.browse_btn.clicked.connect(self.browse_folder)
        self.select_toggle_btn.clicked.connect(self.toggle_select_all)
        self.export_btn.clicked.connect(self.export_excel)
        self.up_btn.clicked.connect(self.move_up)
        self.down_btn.clicked.connect(self.move_down)
        self.convert_btn.clicked.connect(self.convert_checked_files)
        self.pack_btn.clicked.connect(self.generate_pdf)
        self.terminate_btn.clicked.connect(QtWidgets.qApp.quit)
        self.search_box.textChanged.connect(self.apply_filter)
        self.table.cellClicked.connect(self.on_cell_clicked)
        self.table.cellDoubleClicked.connect(self.on_cell_double_clicked)
        self.table.clicked.connect(self.ensure_row_selected)
        self.table.horizontalHeader().sectionClicked.connect(self.handle_header_sort)

        self.timer = QtCore.QTimer(self)
        self.timer.timeout.connect(self.refresh_folder)
        self.timer.start(self.AUTO_REFRESH_INTERVAL_MS)
        self.last_folder = ""
        self.last_fileset = set()

    def toggle_select_all(self):
        indices = self.filtered_indices if self.filtered_indices else list(range(len(self.table_data)))
        for idx in indices:
            self.table_data[idx][0] = self.select_all_state
        self.select_all_state = not self.select_all_state
        self.select_toggle_btn.setText("Deselect All" if self.select_all_state is False else "Select All")
        self.refresh_table()
    def refresh_table(self):
        row = self.table.currentRow()
        self.table.setRowCount(0)
        indices = self.filtered_indices if self.filtered_indices else list(range(len(self.table_data)))
        for ridx, i in enumerate(indices):
            rowdata = self.table_data[i]
            disp_row = list(rowdata)
            disp_row[0] = CHECKED if rowdata[0] else UNCHECKED

            self.table.insertRow(self.table.rowCount())
            base_bg = QtGui.QColor("#f8fbff") if ridx % 2 == 0 else QtGui.QColor("#eaf3ff")
            for col, val in enumerate(disp_row):
                item = QtWidgets.QTableWidgetItem(str(val))
                if col == 0:
                    item.setFont(QtGui.QFont("Times New Roman", 10))
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                    item.setBackground(base_bg)
                elif col == 1:
                    item.setFont(QtGui.QFont("Times New Roman", 10))
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                    ftype = rowdata[1].upper()
                    if ftype == "RTF":
                        item.setBackground(QtGui.QColor("#D6E6FA"))
                    elif ftype == "PDF":
                        item.setBackground(QtGui.QColor("#FFE0BB"))
                    elif ftype == "DOCX":
                        item.setBackground(QtGui.QColor("#C9F6C9"))
                    else:
                        item.setBackground(base_bg)
                elif col == 6:
                    item.setFont(QtGui.QFont("Times New Roman", 10))
                    item.setTextAlignment(QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
                    item.setBackground(base_bg)
                else:
                    item.setFont(QtGui.QFont("Times New Roman", 10))
                    item.setTextAlignment(QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
                    item.setBackground(base_bg)
                self.table.setItem(self.table.rowCount()-1, col, item)
        self.update_button_states()
        if 0 <= row < self.table.rowCount():
            self.table.selectRow(row)

    def handle_header_sort(self, logicalIndex):
        if self.current_sort_column == logicalIndex:
            self.sort_order = (
                QtCore.Qt.DescendingOrder
                if self.sort_order == QtCore.Qt.AscendingOrder
                else QtCore.Qt.AscendingOrder
            )
        else:
            self.sort_order = QtCore.Qt.AscendingOrder
        self.current_sort_column = logicalIndex

        col = logicalIndex
        if col == 0:
            self.table_data.sort(key=lambda x: x[0], reverse=self.sort_order == QtCore.Qt.DescendingOrder)
        elif col == 6:
            self.table_data.sort(key=lambda x: parse_sort_key(x[6]), reverse=self.sort_order == QtCore.Qt.DescendingOrder)
        else:
            self.table_data.sort(key=lambda x: (x[col] or '').lower(), reverse=self.sort_order == QtCore.Qt.DescendingOrder)
        self.apply_filter()
        self.table.horizontalHeader().setSortIndicatorShown(True)
        self.table.horizontalHeader().setSortIndicator(col, self.sort_order)

    def thread_safe_log(self, msg):
        self.log_signal.emit(str(msg))

    def thread_safe_message(self, mtype, msg):
        self.message_signal.emit(mtype, msg)

    def show_messagebox(self, mtype, msg):
        if mtype == "info":
            QtWidgets.QMessageBox.information(self, "Info", msg)
        elif mtype == "error":
            QtWidgets.QMessageBox.critical(self, "Error", msg)
        elif mtype == "done":
            QtWidgets.QMessageBox.information(self, "Done", msg)

    def log(self, msg):
        self.log_box.appendPlainText(str(msg))
        self.log_box.verticalScrollBar().setValue(self.log_box.verticalScrollBar().maximum())
        QtWidgets.QApplication.processEvents()
    def set_button_states(self, browse=True, others=False, convert=False, pack=False):
        self.browse_btn.setEnabled(browse)
        for btn in [self.select_toggle_btn, self.export_btn, self.up_btn, self.down_btn]:
            btn.setEnabled(others)
        self.convert_btn.setEnabled(convert)
        self.pack_btn.setEnabled(pack)
        self.terminate_btn.setEnabled(True)

    def update_button_states(self):
        has_outputs = len(self.table_data) > 0
        any_checked = any(row[0] for row in self.table_data)
        any_rtf_docx_checked = any(row[0] and row[1] in ["RTF", "DOCX"] for row in self.table_data)
        self.set_button_states(
            browse=True,
            others=has_outputs,
            convert=any_rtf_docx_checked and has_outputs,
            pack=any_checked and has_outputs
        )

    def browse_folder(self):
        folder = QtWidgets.QFileDialog.getExistingDirectory(self, "Select Folder")
        if not folder: return
        self.out_folder = folder
        self.last_folder = folder
        self.refresh_folder(full_rescan=True)
        self.log_box.clear()
        self.thread_safe_log("Auto-refresh is enabled (every 3 seconds). Any new file will be auto-detected.")

    def refresh_folder(self, full_rescan=False):
        if not self.out_folder: return
        folder = self.out_folder
        rtf_files = sorted(f for f in os.listdir(folder) if f.lower().endswith('.rtf') and not f.startswith('~$'))
        pdf_files = sorted(f for f in os.listdir(folder) if f.lower().endswith('.pdf') and not f.startswith('~$'))
        docx_files = sorted(f for f in os.listdir(folder) if f.lower().endswith('.docx') and not f.startswith('~$'))

        file_list = [("RTF", f) for f in rtf_files] + [("PDF", f) for f in pdf_files] + [("DOCX", f) for f in docx_files]
        fileset = set(file_list)
        data_lookup = {(row[1], row[2]): i for i, row in enumerate(self.table_data)}
        existing_keys = set(data_lookup.keys())

        # Remove missing files
        i = 0
        while i < len(self.table_data):
            key = (self.table_data[i][1], self.table_data[i][2])
            if key not in fileset:
                del self.table_data[i]
            else:
                i += 1

        # Add new files
        for ftype, fname in file_list:
            if (ftype, fname) not in existing_keys:
                file_path = os.path.join(folder, fname)
                if ftype == 'RTF':
                    t1, t2, t3, bm = extract_titles_from_rtf(file_path)
                elif ftype == 'PDF':
                    t1, t2, t3, bm = extract_title_from_pdf(file_path)
                elif ftype == 'DOCX':
                    t1, t2, t3, bm = extract_title_from_docx(file_path)
                self.table_data.append([True, ftype, fname, t1, t2, t3, bm])
        if not self.sorted_once:
            self.table_data.sort(key=lambda x: parse_sort_key(x[6]))
            self.sorted_once = True
        self.apply_filter()

    def apply_filter(self):
        search = self.search_box.text().strip().lower()
        if not search:
            self.filtered_indices = list(range(len(self.table_data)))
        else:
            self.filtered_indices = []
            for idx, row in enumerate(self.table_data):
                for val in row:
                    if search in str(val).lower():
                        self.filtered_indices.append(idx)
                        break
        self.refresh_table()
    def export_excel(self):
        if not self.table_data or not self.out_folder:
            self.thread_safe_message("error", "No data or folder selected.")
            return
        export_path = os.path.join(self.out_folder, "TLF_Bookmarks.xlsx")
        export_bookmarks_to_excel(self.table_data, export_path)
        self.thread_safe_message("done", f"Bookmarks exported to:\n{export_path}")

    def move_up(self):
        selected = self.table.currentRow()
        if selected <= 0: return
        indices = self.filtered_indices if self.filtered_indices else list(range(len(self.table_data)))
        actual_idx = indices[selected]
        if actual_idx <= 0: return
        self.table_data[actual_idx-1], self.table_data[actual_idx] = self.table_data[actual_idx], self.table_data[actual_idx-1]
        self.apply_filter()
        self.table.selectRow(selected-1)

    def move_down(self):
        selected = self.table.currentRow()
        indices = self.filtered_indices if self.filtered_indices else list(range(len(self.table_data)))
        if selected < 0 or selected >= len(indices) - 1: return
        actual_idx = indices[selected]
        if actual_idx >= len(self.table_data) - 1: return
        self.table_data[actual_idx+1], self.table_data[actual_idx] = self.table_data[actual_idx], self.table_data[actual_idx+1]
        self.apply_filter()
        self.table.selectRow(selected+1)

    def convert_checked_files(self):
        files_to_convert = [row for row in self.table_data if row[0] and row[1] in ['RTF', 'DOCX']]
        if not files_to_convert or not self.out_folder:
            self.thread_safe_message("error", "No checked RTF/DOCX files to convert.")
            return
        rtf_count = sum(1 for row in files_to_convert if row[1] == 'RTF')
        docx_count = sum(1 for row in files_to_convert if row[1] == 'DOCX')
        self.thread_safe_log(f"Total {rtf_count} RTF, {docx_count} DOCX selected for conversion.")
        total = len(files_to_convert)

        def worker():
            try:
                for idx, row in enumerate(files_to_convert):
                    ftype, fname = row[1], row[2]
                    self.thread_safe_log(f"[{idx+1}/{total}] Processing: {fname} ({ftype} to PDF)")
                    file_path = os.path.join(self.out_folder, fname)
                    pdf_path = convert_to_pdf_with_word(file_path, self.out_folder)
                    self.thread_safe_log(f"[Completed] {ftype} ➜ PDF: {os.path.basename(pdf_path)}")
                self.thread_safe_message("done", "Conversion completed for all selected RTF/DOCX files.")
            except Exception as e:
                import traceback
                tb = traceback.format_exc()
                self.thread_safe_log(f"ERROR: {e}\n{tb}")
                self.thread_safe_message("error", f"File Conversion Failed\n\n{e}")

        import threading
        threading.Thread(target=worker, daemon=True).start()

    def generate_pdf(self):
        if not self.table_data or not self.out_folder:
            self.thread_safe_message("error", "No data or folder selected.")
            return
        # === Dynamic PDF Name with versioning ===
        default_prefix = "TLFs_Merged_"
        timestamp_now = datetime.datetime.now().strftime('%Y%m%d_T%H%M')
        suggested_default = f"{default_prefix}{timestamp_now}.pdf"
        user_input = self.outpdf_var.text().strip()

        # If user didn't edit, update timestamp
        if user_input.startswith(default_prefix):
            output_pdf = suggested_default
        else:
            output_pdf = user_input if user_input.lower().endswith(".pdf") else f"{user_input}.pdf"

        # Check if file exists and auto-append _v2, _v3, etc.
        base_name, ext = os.path.splitext(output_pdf)
        counter = 2
        while os.path.exists(os.path.join(self.out_folder, output_pdf)):
            output_pdf = f"{base_name}_v{counter}{ext}"
            counter += 1

        # Set final name in GUI
        self.outpdf_var.setText(output_pdf)
        output_path = os.path.join(self.out_folder, output_pdf)
        files_to_pack = [row for row in self.table_data if row[0]]
        if not files_to_pack:
            self.thread_safe_message("error", "No files selected.")
            return

        rtf_count = sum(1 for row in files_to_pack if row[1] == 'RTF')
        pdf_count = sum(1 for row in files_to_pack if row[1] == 'PDF')
        docx_count = sum(1 for row in files_to_pack if row[1] == 'DOCX')
        self.thread_safe_log(f"Total: {rtf_count} RTF, {pdf_count} PDF, {docx_count} DOCX selected.")

        def worker():
            try:
                temp_pdfs = []
                toc = []
                import fitz
                merged = fitz.open()
                page_count = 0
                for idx, row in enumerate(files_to_pack):
                    ftype, fname, t1, t2, t3, bm = row[1:]
                    self.thread_safe_log(f"[{idx+1}/{len(files_to_pack)}] Adding: {fname} ({ftype})")
                    file_path = os.path.join(self.out_folder, fname)
                    if ftype in ['RTF', 'DOCX']:
                        pdf_path = convert_to_pdf_with_word(file_path, self.out_folder)
                        temp_pdfs.append(pdf_path)
                    else:
                        pdf_path = file_path
                    toc.append([1, bm, page_count + 1])
                    with fitz.open(pdf_path) as doc:
                        merged.insert_pdf(doc)
                        page_count += doc.page_count
                self.thread_safe_log("[Merging] Inserting TOC and bookmarks...")
                temp_bookmarked_pdf = output_path.replace('.pdf', '_bookmarked.pdf')
                merged.set_toc(toc)
                merged.save(temp_bookmarked_pdf)

                if self.toc_checkbox.isChecked():
                    add_toc_and_links(temp_bookmarked_pdf, output_path, font_size=8)
                    os.remove(temp_bookmarked_pdf)
                else:
                    os.rename(temp_bookmarked_pdf, output_path)

                for tfile in temp_pdfs:
                    if os.path.exists(tfile): os.remove(tfile)

                self.thread_safe_log(f"[DONE] PDF created: {output_path}")
                self.thread_safe_message("done", f"Output PDF generated:\n{output_path}")
            except Exception as e:
                import traceback
                tb = traceback.format_exc()
                self.thread_safe_log(f"ERROR: {e}\n{tb}")
                self.thread_safe_message("error", f"PDF Generation Failed\n\n{e}")

        import threading
        threading.Thread(target=worker, daemon=True).start()
    def on_cell_clicked(self, row, col):
        if col == 0:
            indices = self.filtered_indices if self.filtered_indices else list(range(len(self.table_data)))
            actual_idx = indices[row]
            self.table_data[actual_idx][0] = not self.table_data[actual_idx][0]
            self.refresh_table()
        self.table.selectRow(row)

    def on_cell_double_clicked(self, row, col):
        if col == 6:
            indices = self.filtered_indices if self.filtered_indices else list(range(len(self.table_data)))
            actual_idx = indices[row]
            oldval = self.table_data[actual_idx][6]
            newval, ok = QtWidgets.QInputDialog.getText(self, "Edit Bookmark", "Enter new bookmark:", text=oldval)
            if ok and newval is not None:
                self.table_data[actual_idx][6] = newval
                self.refresh_table()
            self.table.selectRow(row)

    def ensure_row_selected(self, index):
        row = index.row()
        self.table.selectRow(row)

# === Application Entry Point ===
def main():
    app = QtWidgets.QApplication(sys.argv)
    window = TLFApp()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
