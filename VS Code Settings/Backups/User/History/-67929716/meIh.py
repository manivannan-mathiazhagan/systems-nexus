# ****************************************************************************************************
# Script Name    : TLF_Packager.py
#
# Purpose        : Scans RTF, DOCX, and PDF files in a folder,
#                  extracts titles (RTF: header/first lines, DOCX: header/body, PDF: page 1),
#                  outputs details to Excel (for user review/order/bookmark edits).
#                  After user approval (optional), combines all into a bookmarked PDF with TOC.
#                  RTF and DOCX are converted to PDF using Microsoft Word (pywin32)
#
# Author         : Manivannan Mathialagan
#
# Created On     : 22-May-2025
# Updated On     : 08-Jul-2025
#
# Parameters
#   1. folder_path             : Path to the folder containing RTF, DOCX, and PDF files.
#   2. output_pdf              : Filename for the output merged PDF.
#   3. delete_converted_pdfs   : Y to delete PDFs converted from RTFs/DOCXs after merging. Default is N
#   4. wait_for_user_approval  : Y to pause for user approval after Excel export (default), N to skip
#
# Example Usage
#   python TLF_Packager.py "input_folder_path" "output_file_name.pdf" Y Y
#
# Notes
#   - Requires Python 3, openpyxl, PyMuPDF, pywin32, python-docx
#   - Requires Microsoft Word (with pywin32) installed
# ****************************************************************************************************




import os
import re
import sys
import fitz  # PyMuPDF for PDF manipulation
import openpyxl # openpyxl for Excel reading/writing 
import docx  # for DOCX support
from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo

def banner(msg):
    print("\n" + "=" * 100)
    print(f" {msg}")
    print("=" * 100)

def print_header():
    print("=" * 100)
    print("TLF Packager - Automated TLF Packaging and Bookmarking Tool")
    print("")
    print("Author: Manivannan Mathialagan")
    print("=" * 100) 

def extract_groups_from_rtf_header(rtf_text):
    header_match = re.search(r'(\\header[lr]?)(.+?)(\\sectd|$)', rtf_text, re.DOTALL)
    if not header_match:
        return []
    header_text = header_match.group(2)
    groups = []
    depth = 0
    buf = ""
    for c in header_text:
        if c == '{':
            if depth == 0:
                buf = ""
            depth += 1
            buf += c
        elif c == '}':
            buf += c
            depth -= 1
            if depth == 0 and buf:
                groups.append(buf)
                buf = ""
        elif depth > 0:
            buf += c
    return groups

def rtf_group_to_text(group):
    group = group.replace('\\cell', '\n').replace('\\row', '\n')
    text = re.sub(r'\\[a-z]+\d*[\s]?|{|}', '', group)
    text = re.sub(r'(cellx\d+|x\d+)', '', text)
    text = re.sub(r'\n+', '\n', text)
    lines = [re.sub(r'^[-: ]+|[-: ]+$', '', l.strip()) for l in text.splitlines() if l.strip()]
    return lines

def extract_header_lines_with_tablepattern(rtf_path):
    try:
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
            rtf_text = f.read()
        groups = extract_groups_from_rtf_header(rtf_text)
        lines = []
        for group in groups:
            lines += rtf_group_to_text(group)
        return lines
    except Exception:
        return []

def valid_title_candidate(candidate):
    c = candidate.strip()
    if not c: return False
    if c.lower().startswith("note:"): return False
    if c.count('.') > 1: return False
    if re.search(r'\[\d+\]', c): return False
    if re.match(r'^\\?\*.*$', c): return False  # Exclude lines starting with * or \*
    if re.match(r'^[A-Z0-9\\\*]+$', c): return False  # Exclude all-caps/keyword lines
    return True

def filter_note_from_line(text):
    return re.split(r'\bNote\s*:', text, flags=re.IGNORECASE)[0].strip()

def extract_titles_from_rtf(rtf_path):
    header_lines = extract_header_lines_with_tablepattern(rtf_path)
    title1, title2, title3 = "", "", ""
    bookmark = ""
    title_pattern = r'(?:(Table|Listing|Figure)\s+)?(Appendix)\s+\d+(?:\.\d+)*[A-Za-z0-9]*|(Table|Listing|Figure)\s+\d+(?:\.\d+)*[A-Za-z0-9]*'
    for i, line in enumerate(header_lines):
        match = re.search(title_pattern, line, re.IGNORECASE)
        if match:
            title1 = match.group(0).strip()
            remainder = line[len(match.group(0)):].strip(" :–-")
            remainder = filter_note_from_line(remainder)
            print(f"\nFile: {os.path.basename(rtf_path)}")
            if remainder:
                valid2 = valid_title_candidate(remainder)
                print(f"Checking Title 2 line: '{remainder}' | Valid: {valid2}")
                if valid2:
                    title2 = remainder
            if not title2 and i+1 < len(header_lines):
                t2 = header_lines[i+1]
                t2_clean = filter_note_from_line(t2)
                valid2 = valid_title_candidate(t2_clean)
                print(f"Checking Title 2 line: '{t2_clean}' | Valid: {valid2}")
                if valid2:
                    title2 = t2_clean.strip()
            if i+2 < len(header_lines):
                t3 = header_lines[i+2]
                t3_clean = filter_note_from_line(t3)
                valid3 = valid_title_candidate(t3_clean)
                print(f"Checking Title 3 line: '{t3_clean}' | Valid: {valid3}")
                if valid3:
                    title3 = t3_clean.strip()
            print("")
            break
    if title1:
        bookmark = title1
        if title2: bookmark += ": " + title2
        if title3: bookmark += " - " + title3
        return title1, title2, title3, bookmark
    return "", "", "", os.path.basename(rtf_path)

def extract_title_from_pdf(pdf_path):
    title_pattern = r'(?:(Table|Listing|Figure)\s+)?(Appendix)\s+\d+(?:\.\d+)*[A-Za-z0-9]*|(Table|Listing|Figure)\s+\d+(?:\.\d+)*[A-Za-z0-9]*'
    try:
        doc = fitz.open(pdf_path)
        text = doc[0].get_text("text")
        doc.close()
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        title1, title2, title3 = "", "", ""
        for i, line in enumerate(lines[:20]):
            match = re.search(title_pattern, line, re.IGNORECASE)
            if match:
                title1 = match.group(0).strip()
                remainder = line[len(match.group(0)):].strip(" :–-")
                remainder = filter_note_from_line(remainder)
                print(f"\nFile: {os.path.basename(pdf_path)} [PDF]")
                if remainder:
                    valid2 = valid_title_candidate(remainder)
                    print(f"Checking Title 2 line: '{remainder}' | Valid: {valid2}")
                    if valid2:
                        title2 = remainder
                if not title2 and i+1 < len(lines):
                    cand2 = lines[i+1]
                    cand2_clean = filter_note_from_line(cand2)
                    valid2 = valid_title_candidate(cand2_clean)
                    print(f"Checking Title 2 line: '{cand2_clean}' | Valid: {valid2}")
                    if valid2:
                        title2 = cand2_clean.strip()
                if i+2 < len(lines):
                    cand3 = lines[i+2]
                    cand3_clean = filter_note_from_line(cand3)
                    valid3 = valid_title_candidate(cand3_clean)
                    print(f"Checking Title 3 line: '{cand3_clean}' | Valid: {valid3}")
                    if valid3:
                        title3 = cand3_clean.strip()
                print("")
                break
        if not title1:
            for idx, line in enumerate(lines[:3]):
                if valid_title_candidate(line):
                    if not title1:
                        title1 = filter_note_from_line(line)
                    elif not title2:
                        title2 = filter_note_from_line(line)
                    elif not title3:
                        title3 = filter_note_from_line(line)
        bookmark = title1
        if title2: bookmark += ": " + title2
        if title3: bookmark += " - " + title3
        if not bookmark.strip():
            bookmark = os.path.basename(pdf_path)
        return title1, title2, title3, bookmark
    except Exception:
        return "", "", "", os.path.basename(pdf_path)

def extract_title_from_docx(docx_path):
    title_pattern = r'(?:(Table|Listing|Figure)\s+)?(Appendix)\s+\d+(?:\.\d+)*[A-Za-z0-9]*|(Table|Listing|Figure)\s+\d+(?:\.\d+)*[A-Za-z0-9]*'
    try:
        doc = docx.Document(docx_path)
        paras = []
        for section in doc.sections:
            for h_para in section.header.paragraphs:
                line = h_para.text.strip()
                if line:
                    paras.append(line)
        print(f"\nFile: {os.path.basename(docx_path)} [DOCX] - Header Preview")
        for i, line in enumerate(paras[:10]):
            print(f"  Header Line {i+1}: '{line}'")
        title1 = title2 = title3 = ""
        for idx, line in enumerate(paras[:10]):
            match = re.match(title_pattern, line, re.IGNORECASE)
            if match:
                title1 = match.group(0).strip()
                remainder = line[len(match.group(0)):].strip(" :–-")
                remainder = filter_note_from_line(remainder)
                print(f"Found Title 1 in header line {idx+1}: '{title1}'")
                if remainder and valid_title_candidate(remainder):
                    title2 = remainder
                    print(f"Found Title 2 (header same line): '{title2}'")
                if not title2 and idx+1 < len(paras):
                    cand2 = paras[idx+1]
                    if valid_title_candidate(cand2):
                        title2 = cand2
                        print(f"Found Title 2 (header next line): '{title2}'")
                if idx+2 < len(paras):
                    cand3 = paras[idx+2]
                    if valid_title_candidate(cand3):
                        title3 = cand3
                        print(f"Found Title 3 (header): '{title3}'")
                break
        if not title1:
            body_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            body_paras.append(cell_text)
            print(f"\nFile: {os.path.basename(docx_path)} [DOCX] - Body/Table Preview")
            for i, line in enumerate(body_paras[:10]):
                print(f"  Body Line {i+1}: '{line}'")
            for idx, line in enumerate(body_paras[:15]):
                match = re.match(title_pattern, line, re.IGNORECASE)
                if match:
                    title1 = match.group().strip()
                    remainder = line[len(match.group(0)):].strip(" :–-")
                    remainder = filter_note_from_line(remainder)
                    print(f"Found Title 1 in body line {idx+1}: '{title1}'")
                    if remainder and valid_title_candidate(remainder):
                        title2 = remainder
                        print(f"Found Title 2 (body same line): '{title2}'")
                    if not title2 and idx+1 < len(body_paras):
                        cand2 = body_paras[idx+1]
                        if valid_title_candidate(cand2):
                            title2 = cand2
                            print(f"Found Title 2 (body next line): '{title2}'")
                    if idx+2 < len(body_paras):
                        cand3 = body_paras[idx+2]
                        if valid_title_candidate(cand3):
                            title3 = cand3
                            print(f"Found Title 3 (body): '{title3}'")
                    break
            if not title1:
                print(f"\nFile: {os.path.basename(docx_path)} [DOCX] - Fallback mode")
                all_lines = paras + body_paras
                found_lines = 0
                for line in all_lines[:15]:
                    if valid_title_candidate(line):
                        found_lines += 1
                        if not title1:
                            print(f"Fallback Title 1: '{line}'")
                            title1 = filter_note_from_line(line)
                        elif not title2:
                            print(f"Fallback Title 2: '{line}'")
                            title2 = filter_note_from_line(line)
                        elif not title3:
                            print(f"Fallback Title 3: '{line}'")
                            title3 = filter_note_from_line(line)
                        if found_lines == 3:
                            break
        bookmark = title1
        if title2: bookmark += ": " + title2
        if title3: bookmark += " - " + title3
        if not bookmark.strip():
            bookmark = os.path.basename(docx_path)
        return title1, title2, title3, bookmark
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return "", "", "", os.path.basename(docx_path)

def parse_sort_key(bookmark):
    order = {"table": 0, "listing": 1, "figure": 2}
    text = (bookmark or "").lower()
    text = re.sub(r'^[^a-zA-Z]*', '', text)
    if 'appendix' in text:
        match = re.search(r'appendix\s*(\d+)', text)
        num = [int(match.group(1))] if match else [9999]
        return (3, num)
    for key in order:
        if key in text:
            match = re.search(r'(\d+(\.\d+)*)', text)
            num = [int(x) for x in match.group(1).split('.')] if match else [0]
            return (order[key], num)
    return (99, [0])

def write_titles_to_excel(entries, excel_file):
    sorted_entries = sorted(entries, key=lambda x: parse_sort_key(x['bookmark']))
    wb = Workbook()
    ws = wb.active
    headers = ["SourceType", "Filename", "Title 1", "Title 2", "Title 3", "Bookmark"]
    for i, header in enumerate(headers, start=1):
        ws.cell(row=1, column=i, value=header)
    for row_idx, entry in enumerate(sorted_entries, start=2):
        ws.cell(row=row_idx, column=1, value=entry['type'])
        ws.cell(row=row_idx, column=2, value=os.path.basename(entry['file']))
        ws.cell(row=row_idx, column=3, value=entry['title1'])
        ws.cell(row=row_idx, column=4, value=entry['title2'])
        ws.cell(row=row_idx, column=5, value=entry['title3'])
        ws.cell(row=row_idx, column=6, value=entry['bookmark'])
    column_widths = [15, 40, 35, 40, 30, 60]
    for idx, col_width in enumerate(column_widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(idx)].width = col_width
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical='top')
    last_row = ws.max_row
    table_ref = f"A1:F{last_row}"
    tab = Table(displayName="TLFBookmarks", ref=table_ref)
    style = TableStyleInfo(name="TableStyleMedium9", showRowStripes=True, showColumnStripes=False)
    tab.tableStyleInfo = style
    ws.add_table(tab)
    ws.freeze_panes = "A2"
    wb.save(excel_file)
    print(f"\n[Step 1] Titles and bookmarks exported to Excel for review: {excel_file}\n"
          "Instructions:\n"
          " - Open the Excel file.\n"
          " - Reorder rows and/or edit the 'Bookmark' column for PDF navigation as required.\n"
          " - SAVE and CLOSE the file before continuing.")

def gui_pause_and_debug(excel_file):
    import openpyxl
    def print_excel_titles():
        wb = openpyxl.load_workbook(excel_file)
        ws = wb.active
        print("\n[Excel Review] Extracted Titles and Bookmarks (from Excel):")
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row and any(row):
                print(f"  [{row[0]}] {row[1]}: {row[5]}")
        wb.close()
    try:
        import tkinter as tk
        from tkinter import messagebox
        root = tk.Tk()
        root.withdraw()
        result = messagebox.askquestion(
            "TLF Packager",
            "Awaiting user approval of Excel.\n\nClick 'Yes' to generate the PDF, or 'No' to stop and print titles/bookmarks."
        )
        root.destroy()
        if result != "yes":
            print("\n[STOP] PDF generation stopped by user after Excel review.\n")
            print_excel_titles()
            print("\nPlease review your Excel file and run the script again when ready.")
            sys.exit(0)
    except Exception:
        try:
            answer = input("Awaiting user approval of Excel. (Type 'y' to continue, any other key to stop and print titles): ")
            if answer.lower() != "y":
                print("\n[STOP] PDF generation stopped by user after Excel review.\n")
                print_excel_titles()
                print("\nPlease review your Excel file and run the script again when ready.")
                sys.exit(0)
        except:
            sys.exit(0)

def read_entries_from_excel(excel_file, folder):
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active
    entries = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        src_type, filename, title1, title2, title3, bookmark = row
        filepath = os.path.join(folder, filename)
        entries.append({
            'type': src_type,
            'file': filepath,
            'title1': title1,
            'title2': title2,
            'title3': title3,
            'bookmark': bookmark
        })
    return entries

def convert_to_pdf_with_word(input_path, output_folder):
    import win32com.client
    import time
    pdf_path = os.path.splitext(input_path)[0] + ".pdf"
    pdf_path = os.path.join(output_folder, os.path.basename(pdf_path))
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(input_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close(False)
        time.sleep(0.5)
    except Exception as e:
        print(f"[Word] Error converting {input_path}: {e}")
    finally:
        word.Quit()
    print(f"[Word Converted] {os.path.basename(input_path)}")
    return pdf_path

def merge_pdfs_with_bookmarks(entries, output_path, delete_intermediate=False):
    banner("STEP 4: Bookmark Creation, Merging, and PDF Generation")
    merged = fitz.open()
    toc = []
    page_count = 0
    to_delete = []
    print("Merging the following files (in order):")
    for idx, entry in enumerate(entries, 1):
        if entry['type'] in ['RTF', 'DOCX']:
            pdf_path = convert_to_pdf_with_word(entry['file'], os.path.dirname(entry['file']))
            to_delete.append(pdf_path)
        else:
            pdf_path = entry['file']
        print(f"  {idx:2d}. [{entry['type']}] {os.path.basename(pdf_path)}  -->  {entry['bookmark']}")
        with fitz.open(pdf_path) as doc:
            bookmark = (entry['bookmark'] or os.path.basename(entry['file'])).strip()
            toc.append([1, bookmark, page_count + 1])
            merged.insert_pdf(doc)
            page_count += doc.page_count
    merged.set_toc(toc)
    merged.save(output_path)
    print(f"\n[Step 4] Final merged PDF with bookmarks: {output_path}")
    if delete_intermediate:
        print("\nCleaning up PDFs generated from RTFs/DOCXs:")
        for pdf in to_delete:
            if os.path.exists(pdf):
                os.remove(pdf)
                print(f"  [Deleted] {os.path.basename(pdf)}")

def wrap_toc_title(title, font_size, max_width):
    words = (title or "").split()
    lines = []
    current_line = ""
    for word in words:
        test_line = current_line + (" " if current_line else "") + word
        if fitz.get_text_length(test_line, fontsize=font_size) <= max_width:
            current_line = test_line
        else:
            if current_line:
                lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return [line for line in lines if line.strip()]

def add_toc_and_links(bookmarked_pdf, final_pdf, font_size=8):
    banner("STEP 5: Table of Contents (TOC) Generation")
    def extract_toc_entries(doc):
        return [(lvl, title.strip(), page_num - 1) for lvl, title, page_num in doc.get_toc(simple=True)]
    def paginate_wrapped_entries(toc_entries, font_size, toc_text_max_width, page_height):
        y_spacing = font_size * 1.5
        lines_per_page = int((page_height - 100) // y_spacing)
        lines_per_page -= 2
        paginated_entries, current_page_entries, current_line_count = [], [], 0
        for entry in toc_entries:
            level, title, target_page = entry
            indent = 20 * (level - 1)
            entry_max_width = toc_text_max_width - indent
            wrapped_lines = wrap_toc_title(title, font_size, entry_max_width)
            line_count = len(wrapped_lines)
            if current_line_count + line_count > lines_per_page:
                paginated_entries.append(current_page_entries)
                current_page_entries, current_line_count = [], 0
            current_page_entries.append((entry, wrapped_lines))
            current_line_count += line_count
        if current_page_entries:
            paginated_entries.append(current_page_entries)
        return paginated_entries

    def generate_toc_pages(paginated_entries, font_size, page_width, page_height, toc_page_count, toc_text_max_width, page_number_x, page_number_width, gap, right_margin):
        toc_doc = fitz.open()
        link_targets = []
        left_margin, top_margin = 50, 50
        y_spacing = font_size * 1.5
        for page_index, entries in enumerate(paginated_entries):
            page = toc_doc.new_page(width=page_width, height=page_height)
            y = top_margin
            if page_index == 0:
                toc_title = "Table of Contents"
                title_width = fitz.get_text_length(toc_title, fontsize=font_size + 2)
                page.insert_text(
                    (page_width / 2 - title_width / 2, y),
                    toc_title,
                    fontsize=font_size + 2,
                    fontname="Helvetica",
                    color=(0, 0, 0)
                )
                y += font_size * 2
            for (level, title, target_page), wrapped_lines in entries:
                indent = 20 * (level - 1)
                x = left_margin + indent
                page_number_str = str(target_page + toc_page_count + 1)
                page_number_width_local = fitz.get_text_length(page_number_str, fontsize=font_size)
                first_line_y = y
                for i, line in enumerate(wrapped_lines):
                    line_width = fitz.get_text_length(line, fontsize=font_size)
                    if i == len(wrapped_lines) - 1:
                        page.insert_text((x, y), line, fontsize=font_size)
                        dots_start_x = x + line_width + 2
                        if dots_start_x < page_number_x - page_number_width_local - gap:
                            dots_width = page_number_x - page_number_width_local - gap - dots_start_x
                            dot_char_width = fitz.get_text_length('.', fontsize=font_size)
                            dot_count = int(dots_width // dot_char_width)
                            dots_str = '.' * dot_count
                            page.insert_text((dots_start_x, y), dots_str, fontsize=font_size)
                        page.insert_text(
                            (page_number_x - page_number_width_local, y),
                            page_number_str, fontsize=font_size
                        )
                    else:
                        page.insert_text((x, y), line, fontsize=font_size)
                    y += y_spacing
                rect = fitz.Rect(x, first_line_y - font_size, page_number_x, y)
                link_targets.append((page_index, rect, target_page))
        return toc_doc, link_targets

    base = fitz.open(bookmarked_pdf)
    original_toc = extract_toc_entries(base)
    original_bookmarks = base.get_toc()
    width, height = fitz.paper_size("a4")
    left_margin, right_margin = 50, 60
    max_page_number = 99999
    page_number_width = fitz.get_text_length(str(max_page_number), fontsize=font_size)
    page_number_x = width - right_margin
    gap = 8
    toc_text_max_width = page_number_x - left_margin - page_number_width - gap

    paginated = paginate_wrapped_entries(original_toc, font_size, toc_text_max_width, height)
    toc_doc, link_targets = generate_toc_pages(
        paginated, font_size, width, height, len(paginated),
        toc_text_max_width, page_number_x, page_number_width, gap, right_margin
    )
    toc_page_count = len(paginated)

    final = fitz.open()
    final.insert_pdf(toc_doc)
    final.insert_pdf(base)

    def add_toc_hyperlinks(doc, link_targets, toc_page_count):
        for toc_page_index, rect, target_page in link_targets:
            doc[toc_page_index].insert_link({
                "kind": fitz.LINK_GOTO,
                "from": rect,
                "page": target_page + toc_page_count
            })
    add_toc_hyperlinks(final, link_targets, toc_page_count)

    def shift_bookmark_pages(bookmarks, offset):
        return [[lvl, title, page + offset] for lvl, title, page in bookmarks if len([lvl, title, page]) >= 3]
    def add_existing_bookmarks(doc, bookmarks, offset):
        doc.set_toc(shift_bookmark_pages(bookmarks, offset))
    add_existing_bookmarks(final, original_bookmarks, toc_page_count)

    final.save(final_pdf)
    print(f"\n[Step 5] Final PDF with TOC created: {final_pdf}")
    final.close()
    base.close()
    try:
        os.remove(bookmarked_pdf)
        print(f"\n[Deleted intermediate] {os.path.basename(bookmarked_pdf)}")
    except Exception as e:
        print(f"[Warning] Couldn't delete intermediate: {e}")

def main(folder, output_pdf_name="TLFs.pdf", delete_intermediate_pdfs=False, wait_for_user_approval=True):
    print_header()
    banner("STEP 1: Detecting Input Files and Exporting for Excel Review")
    print(f"Input folder: {folder}")

    rtf_files = sorted(f for f in os.listdir(folder) if f.lower().endswith('.rtf') and not f.startswith('~$'))
    pdf_files = sorted(f for f in os.listdir(folder) if f.lower().endswith('.pdf') and not f.startswith('~$'))
    docx_files = sorted(f for f in os.listdir(folder) if f.lower().endswith('.docx') and not f.startswith('~$'))

    print(f"Found {len(rtf_files)} RTF files, {len(pdf_files)} PDF files, and {len(docx_files)} DOCX files.")

    print("\n--- Listing all RTF files below ---")
    for f in rtf_files:
        print(f"   - {f}")

    print("\n--- Listing all PDF files below ---")
    for f in pdf_files:
        print(f"   - {f}")

    print("\n--- Listing all DOCX files below ---")
    for f in docx_files:
        print(f"   - {f}")

    print("\n\nTitle checking starts\n")
    entries = []
    for rtf in rtf_files:
        rtf_path = os.path.join(folder, rtf)
        title1, title2, title3, bookmark = extract_titles_from_rtf(rtf_path)
        entries.append({'type': 'RTF', 'file': rtf_path, 'title1': title1, 'title2': title2, 'title3': title3, 'bookmark': bookmark})
    for pdf in pdf_files:
        pdf_path = os.path.join(folder, pdf)
        title1, title2, title3, bookmark = extract_title_from_pdf(pdf_path)
        entries.append({'type': 'PDF', 'file': pdf_path, 'title1': title1, 'title2': title2, 'title3': title3, 'bookmark': bookmark})
    for docx_file in docx_files:
        docx_path = os.path.join(folder, docx_file)
        title1, title2, title3, bookmark = extract_title_from_docx(docx_path)
        entries.append({'type': 'DOCX', 'file': docx_path, 'title1': title1, 'title2': title2, 'title3': title3, 'bookmark': bookmark})
    print("\nTitle checking ends\n")

    excel_file = os.path.join(folder, "TLF_Bookmark_Worksheet.xlsx")
    write_titles_to_excel(entries, excel_file)

    banner("STEP 2: User Review and Approval in Excel")
    if wait_for_user_approval:
        print("User action required:")
        print(f"1. Open the generated Excel file ({os.path.basename(excel_file)}).")
        print("2. Reorder rows and/or edit 'Bookmark' as desired for correct PDF navigation.")
        print("3. SAVE and CLOSE the file before continuing.")
        gui_pause_and_debug(excel_file)
    else:
        print("Skipping user approval as per input parameter. Proceeding directly.")

    banner("STEP 3: Reading Approved Bookmarks and Order from Excel")
    entries = read_entries_from_excel(excel_file, folder)
    print(f"Will merge {len(entries)} entries in this order:")
    for i, e in enumerate(entries, 1):
        print(f"  {i:2d}. [{e['type']}] {os.path.basename(e['file'])} --> {e['bookmark']}")

    bookmarked = os.path.join(folder, "TLFs_bookmarked.pdf")
    final = os.path.join(folder, output_pdf_name)
    merge_pdfs_with_bookmarks(entries, bookmarked, delete_intermediate=delete_intermediate_pdfs)
    add_toc_and_links(bookmarked, final, font_size=8)

    banner("SUMMARY")
    print(f"Completed packaging for folder: {folder}")
    print(f"Output PDF: {final}")
    print("Process complete.")

if __name__ == "__main__":
    if len(sys.argv) < 3 or len(sys.argv) > 5:
        print("Usage: python TLF_Packager.py <folder_path> <output_pdf_name> [delete_converted_pdfs: Y/N] [wait_for_user_approval: Y/N]")
        sys.exit(1)

    path = sys.argv[1]
    output_pdf_name = os.path.basename(sys.argv[2])
    delete_intermediate_pdfs = sys.argv[3].strip().upper() == "Y" if len(sys.argv) > 3 else False
    wait_for_user_approval = sys.argv[4].strip().upper() == "Y" if len(sys.argv) > 4 else True

    if not os.path.isdir(path):
        print(f"Error: Folder not found - {path}")
        sys.exit(1)

    main(path, output_pdf_name, delete_intermediate_pdfs, wait_for_user_approval)
